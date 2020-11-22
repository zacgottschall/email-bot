import urllib, json, requests
from datetime import datetime
from titlecase import titlecase
from xml.etree import ElementTree
from progress.bar import Bar
import docx
from .add_hyperlink import add_hyperlink
from pathlib import Path
from ..task import Task
from ...models import OutgoingEmail


class FederalRegisterTask(Task):

    def do_task(self, start_date, end_date) -> OutgoingEmail:
        date_of_report = end_date
        document, filename = initialize_document(date_of_report)
        website_url, json_url = get_urls(start_date, end_date)
        results, total, total_pages = fetch_search_results(json_url)

        count, currentPage = 0, 1
        with Bar('Generating report...', max=total) as bar:
            while currentPage <= total_pages:
                for notice in results:
                    title, document_number, agencies, publication_date, pdf_url = destructure_notice(notice)
                    start_page, end_page, volume, issue, _type = fetch_more_info(document_number)

                    add_notice_entry(
                        document, title, document_number, agencies, publication_date,
                        start_page, end_page, volume, issue, _type, pdf_url
                    )

                    count += 1
                    bar.next()

                currentPage += 1
                results = fetch_search_results(json_url, currentPage)

        document.save(filename)

        subject = "Federal Register Report - {}".format(date_of_report)
        body = "i love u"
        return OutgoingEmail(subject=subject, body=body, filename=filename)


def initialize_document(date_of_report):
    document = docx.Document()
    style = document.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = docx.shared.Pt(12)
    filename = '/email-bot/' + 'FR_REPORT_' + date_of_report.replace('/', '-') + '.docx'
    return document, filename


def get_urls(start_date, end_date):
    search_conditions = "conditions%5Bpublication_date%5D%5Bgte%5D=" + \
                        urllib.parse.quote(start_date, safe='') + "&conditions%5Bpublication_date%5D%5Blte%5D=" + \
                        urllib.parse.quote(end_date,
                                           safe='') + "&conditions%5Bterm%5D=Tribal+%7C+Indian+%7C+Tribe+%7C+%22Native+American%22+%7C+%22Alaska+Native%22&order=oldest"

    website_url = "https://www.federalregister.gov/documents/search?" + search_conditions
    json_url = "https://www.federalregister.gov/api/v1/documents.json?" + search_conditions

    return website_url, json_url


def fetch_search_results(search_url, page=1):
    if page > 1:
        search_url += '&page=' + str(page)

    with urllib.request.urlopen(search_url) as url:
        response = json.loads(url.read().decode())

        if page == 1:
            return response['results'], response['count'], response['total_pages']
        else:
            return response['results']


def destructure_notice(notice):
    publication_date = datetime.strptime(notice['publication_date'], '%Y-%m-%d').strftime('%B %d, %Y')
    title = notice['title']
    agencies = ", ".join((titlecase(a['raw_name']) for a in notice['agencies']))
    document_number = notice['document_number']
    pdf_url = notice['pdf_url']
    return title, document_number, agencies, publication_date, pdf_url


def fetch_more_info(document_number):
    more_info_url = 'https://www.federalregister.gov/api/v1/documents/' + document_number + '.json?fields%5B%5D=start_page&fields%5B%5D=end_page&fields%5B%5D=volume&fields%5B%5D=type&fields%5B%5D=mods_url'
    with urllib.request.urlopen(more_info_url) as url:
        more_info_response = json.loads(url.read().decode())
        start_page = str(more_info_response['start_page'])
        end_page = str(more_info_response['end_page'])
        volume = str(more_info_response['volume'])
        _type = more_info_response['type']
        mods_url = more_info_response['mods_url']

    mods_request = requests.get(mods_url)
    tree = ElementTree.fromstring(mods_request.content)

    for n in tree.iter('{http://www.loc.gov/mods/v3}issue'):
        issue = n.text

    return start_page, end_page, volume, issue, _type


def add_notice_entry(document, title, document_number, agencies, publication_date, start_page, end_page, volume, issue,
                     _type, pdf_url):
    p = document.add_paragraph(style='List Number')
    hyperlink = add_hyperlink(p, pdf_url, agencies + ': ' + title, '002776', False)
    p.add_run(
        ' [Federal Register: Volume ' + volume + ', Number ' + issue + ' (' + publication_date + ')][' + _type + '][Pages ' + start_page + '-' + end_page + '].')
